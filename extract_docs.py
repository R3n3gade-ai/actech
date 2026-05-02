"""Extract canonical ARMS spec docs to plain text for grep/audit."""
import os, sys
from pathlib import Path
from docx import Document

src_dir = Path("data/Original documents")
out_dir = Path("data/_doc_extracts")
out_dir.mkdir(exist_ok=True)

files = sorted(src_dir.glob("*.docx"))
print(f"Extracting {len(files)} docx files -> {out_dir}/")

for f in files:
    try:
        doc = Document(f)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        # Tables
        for ti, tbl in enumerate(doc.tables):
            lines.append(f"\n[TABLE {ti}]")
            for row in tbl.rows:
                cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
                lines.append(" || ".join(cells))
        out = out_dir / (f.stem + ".txt")
        out.write_text("\n".join(lines), encoding="utf-8")
        print(f"  {f.name} -> {out.name} ({len(lines)} lines)")
    except Exception as e:
        print(f"  FAILED {f.name}: {e}")
